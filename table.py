import json
import streamlit as st
from docx import Document
from io import BytesIO
import base64

# Function to create the question tables document
def create_question_tables(data):
    doc = Document()

    for question in data:
        # Create a table with 2 columns for metadata
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'

        # Add metadata rows
        rows = [
            ("LU", question.get("LU", "N/A")),
            ("LUID", question.get("LUID", "N/A")),
            ("LO", question.get("LO", "N/A")),
            ("LOID", question.get("LOID", "N/A")),
            ("Question Type", question.get("Question Type", "N/A")),
            ("Difficulty Level", question.get("Difficulty Level", "N/A")),
            ("Cognitive Dimension (Bloom's Level)", question.get("Cognitive Dimension (Bloom’s Level)", "N/A")),
            ("Question ID", question.get("Question ID", "N/A")),
            ("Title (Stem and Prompt)", question.get("Title (Stem and Prompt)", "N/A")),
        ]

        for category, detail in rows:
            row = table.add_row()
            row.cells[0].text = category
            row.cells[1].text = detail

        # Add a separate 4-column section for options
        options_table = doc.add_table(rows=1, cols=4)
        options_table.style = 'Table Grid'

        # Add headers for the options table
        hdr_cells = options_table.rows[0].cells
        hdr_cells[0].text = "Option"
        hdr_cells[1].text = "Text"
        hdr_cells[2].text = "Type of Distractor"
        hdr_cells[3].text = "Distractor Rationale"

        # Add options rows
        for option in question.get("Options", []):
            row = options_table.add_row()
            row.cells[0].text = option.get("Option", "N/A")
            row.cells[1].text = option.get("Text", "N/A")
            row.cells[2].text = option.get("Type of Distractor", "N/A")
            row.cells[3].text = option.get("Distractor Rationale", "N/A")

        # Add a new table for Hint and Solution
        bottom_table = doc.add_table(rows=0, cols=2)
        bottom_table.style = 'Table Grid'

        # Add bottom rows
        bottom_rows = [
            ("Hint", question.get("Hint", "N/A")),
            ("Solution", question.get("Solution", "N/A"))
        ]

        for category, detail in bottom_rows:
            row = bottom_table.add_row()
            row.cells[0].text = category
            row.cells[1].text = detail

        # Add a page break after each question
        doc.add_page_break()

    # Save the document to a BytesIO object
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# Function to create a download link
def create_download_link(file_stream, filename):
    b64 = base64.b64encode(file_stream.read()).decode()
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="{filename}">Download the generated document</a>'
    return href

# Streamlit app layout
def main():
    st.title("Question Table Generator")
    st.markdown("### Paste the JSON data below to generate the Google Doc-style output")

    # JSON input from user
    json_input = st.text_area("Enter JSON Data:", height=350)

    if st.button("Generate Document"):
        try:
            # Parse JSON input
            data = json.loads(json_input)

            # Validate data structure
            if not isinstance(data, list):
                raise ValueError("Input JSON must be a list of questions.")
            for question in data:
                if "Options" not in question or not isinstance(question["Options"], list):
                    raise ValueError("Each question must contain an 'Options' field with a list of options.")

            # Create document
            file_stream = create_question_tables(data)

            # Create download link
            st.success("Document generated successfully!")
            download_link = create_download_link(file_stream, "Generated_Questions.docx")
            st.markdown(download_link, unsafe_allow_html=True)
        except json.JSONDecodeError:
            st.error("Invalid JSON format. Please check your input.")
        except Exception as e:
            st.error(f"Error: {str(e)}")

if __name__ == "__main__":
    main()
